import pdfplumber
from PIL import Image
import pytesseract
from docx import Document
import os

os.environ["TESSDATA_PREFIX"] = "tessdata/"  

def extract_pdf_text_with_ocr_to_doc(pdf_path, doc_path):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            # Convert PDF page to image for OCR
            page_image = page.to_image(resolution=300)
            pil_img = page_image.original
            ocr_text = pytesseract.image_to_string(pil_img, lang="eng+mal") 
            doc.add_heading(f"Page {i}", level=2)
            doc.add_paragraph(ocr_text if ocr_text else "[No text found]")
    doc.save(doc_path)
    print(f"OCR extracted text saved to {doc_path}")

if __name__ == "__main__":
    pdf_path = "PDF-BOOK.pdf"  # Replace with your PDF file path
    doc_path = "extracted_text-part2.docx"  # Output DOC file
    extract_pdf_text_with_ocr_to_doc(pdf_path, doc_path)