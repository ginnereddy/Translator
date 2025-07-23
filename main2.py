# use amazon textract for extracting the text from the pdf
import json
import base64
import pdfplumber
from docx import Document
import re

import boto3
translate = boto3.client(service_name='translate', region_name='us-east-2', use_ssl=True)

def split_doc_by_pages(doc):
    """
    Splits the document into pages based on 'Page X' markers.
    Returns a list of tuples: (page_title, page_text)
    """
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Split based on "Page X" markers (e.g., Page 1, Page 2)
    page_splits = re.split(r'(Page\s+\d+)', full_text) 
    
    pages = []
    for i in range(1, len(page_splits), 2):
        page_title = page_splits[i].strip()
        page_content = page_splits[i + 1].strip() if (i + 1) < len(page_splits) else ''
        pages.append((page_title, page_content))
    
    return pages
def translate_text_block(text, source_language='ml', target_language='en'):
    """
    Translates a block of text using AWS Translate.
    """
    if not text.strip():
        return ""
    response = translate.translate_text(
        Text=text,
        SourceLanguageCode=source_language,
        TargetLanguageCode=target_language
    )
    return response.get('TranslatedText', '')


def translate_docx_by_pages(input_docx, output_docx, source_language='ml', target_language='en'):
    doc = Document(input_docx)
    translated_doc = Document()

    pages = split_doc_by_pages(doc)

    for page_title, page_text in pages:
        # Add original page title as bold and font size 14
        title_para = translated_doc.add_paragraph()
        run = title_para.add_run(page_title)
        run.bold = True
        run.font.size = 140000  # 14 pt in docx units

        # Translate content
        translated_text = translate_text_block(page_text, source_language, target_language)
        
        # Add translated content
        translated_doc.add_paragraph(translated_text)
        
        # Optional: Add a page break or spacing
        translated_doc.add_paragraph("\n")

    translated_doc.save(output_docx)
    print(f"Translation complete. Saved to {output_docx}")

if __name__ == "__main__":
    input_docx = "extracted_text-part2.docx"  # Input DOCX with extracted text
    output_docx = "translated_text-part2.docx"  # Output DOCX for translated text
    translate_docx_by_pages(input_docx, output_docx)
